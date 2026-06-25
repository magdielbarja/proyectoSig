import os
import pandas as pd
from docx import Document

def extract_docx(docx_path, txt_path):
    print(f"Extracting {docx_path} -> {txt_path}")
    try:
        doc = Document(docx_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                fullText.append(" | ".join(row_text))
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(fullText))
        print("Success")
    except Exception as e:
        print(f"Error: {e}")

def inspect_xls(xls_path, out_path):
    print(f"Inspecting {xls_path} -> {out_path}")
    try:
        # Load XLS
        xl = pd.ExcelFile(xls_path)
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(f"Sheets in file: {xl.sheet_names}\n\n")
            for sheet in xl.sheet_names:
                df = xl.parse(sheet)
                f.write(f"=== Sheet: {sheet} ===\n")
                f.write(f"Shape: {df.shape}\n")
                f.write(f"Columns: {list(df.columns)}\n\n")
                f.write("First 10 rows:\n")
                f.write(df.head(10).to_string())
                f.write("\n\n" + "="*50 + "\n\n")
        print("Success")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    docs_dir = r"c:\Users\migue\Desktop\sig\proyecto\Proyecto_1_2026\Documentos"
    data_dir = r"c:\Users\migue\Desktop\sig\proyecto\Proyecto_1_2026\Datos_Lineas"
    
    extract_docx(
        os.path.join(docs_dir, "Alcance del Proyecto microbuses SIG 2_2025 final VER2.docx"),
        os.path.join(docs_dir, "Alcance_del_Proyecto.txt")
    )
    extract_docx(
        os.path.join(docs_dir, "Como preparar los datos de la red de microbus.docx"),
        os.path.join(docs_dir, "Como_preparar_los_datos.txt")
    )
    inspect_xls(
        os.path.join(data_dir, "DatosLineas.xls"),
        os.path.join(data_dir, "DatosLineas_summary.txt")
    )
